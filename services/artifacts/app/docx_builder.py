"""DocContent -> .docx bytes (python-docx)."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from . import brand
from .models import DocContent


def _set_font(run, name: str, size: int, color: str, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _render_sections(doc, sections):
    for s in sections:
        doc.add_heading(s.heading, level=2)
        for b in s.blocks:
            if b.type == "paragraph":
                doc.add_paragraph(b.text)
            elif b.type in ("bullets", "numbers"):
                style = "List Bullet" if b.type == "bullets" else "List Number"
                for it in b.items:
                    doc.add_paragraph(it, style=style)
            elif b.type == "callout":
                doc.add_paragraph(f"{b.value} — {b.label}")
            elif b.type == "table":
                t = doc.add_table(rows=1, cols=len(b.columns))
                t.style = "Light Grid Accent 1"
                for i, col in enumerate(b.columns):
                    t.rows[0].cells[i].text = col
                for row in b.rows:
                    cells = t.add_row().cells
                    for i in range(len(b.columns)):
                        cells[i].text = row[i] if i < len(row) else ""
            # 'bars' blocks are omitted from docx in v1 (best-effort text export)


def build_doc(content: DocContent, name: str) -> bytes:
    doc = Document()
    doc.core_properties.title = name
    normal = doc.styles["Normal"]
    normal.font.name = brand.UI_FONT
    normal.font.size = Pt(11)

    eyebrow = doc.add_paragraph()
    _set_font(eyebrow.add_run(content.eyebrow.upper()), brand.UI_FONT, 8, brand.MUTED, bold=True)

    title = doc.add_paragraph()
    _set_font(title.add_run(content.title), brand.SERIF_FONT, 26, brand.INK, bold=True)

    meta = doc.add_paragraph()
    _set_font(meta.add_run(content.meta), brand.UI_FONT, 9, brand.MUTED)

    doc.add_paragraph()  # spacer

    if content.sections:
        _render_sections(doc, content.sections)
    else:
        for p in content.paragraphs or []:
            para = doc.add_paragraph()
            _set_font(para.add_run(p), brand.UI_FONT, 11, brand.INK)

    if content.bars:
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ("Period", "Index")):
            cell.paragraphs[0].add_run(label).bold = True
        for b in content.bars:
            row = table.add_row().cells
            row[0].text = b.label
            row[1].text = f"{round(b.value * 100)}"

    if content.callout:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(c.add_run(f"{content.callout.value}  "), brand.SERIF_FONT, 20, brand.POSITIVE, bold=True)
        _set_font(c.add_run(content.callout.label), brand.UI_FONT, 9, brand.MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
